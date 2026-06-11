from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "outputs" / "ctrl1_qec_analysis_2026-04-25"
DOCX_PATH = OUT_DIR / "ctrl1_qec_run_main_takeaways.docx"


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        run.font.size = Pt(10)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(20)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(13)
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("CTRL-1 Qec SAC Run: Main Takeaways")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run("Brief notes for results and discussion")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(90, 90, 90)

    snapshot = doc.add_paragraph()
    snapshot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    snapshot.paragraph_format.space_after = Pt(12)
    r = snapshot.add_run(
        "Run snapshot: 58,463 steps | mean reward 0.3995 | mean J ratio 0.9500 | final reward 0.7755"
    )
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(31, 78, 121)

    doc.add_heading("Best takeaways", level=1)
    add_bullets(
        doc,
        [
            "The run is stable and usable: no missing rewards, no Qec bound violations, and no SNO_3 below 0.5.",
            "CTRL-1 improved the mean proxy objective versus manual baseline: mean J ratio 0.9500 vs 1.0332.",
            "Final behavior was encouraging: final reward 0.7755 and final ratio 0.8449.",
        ],
    )

    doc.add_heading("Main weaknesses", level=1)
    add_bullets(
        doc,
        [
            "Mean reward was lower than manual baseline: 0.3995 vs 0.6611.",
            "About 17.2% of rows were clipped at reward = -1, showing unstable or poor local periods.",
            "The SAC phase did not show a clean monotonic learning improvement; last 5,000-step mean reward was only 0.3757.",
        ],
    )

    doc.add_heading("Discussion angle", level=1)
    add_bullets(
        doc,
        [
            "CTRL-1 is promising as a cost-oriented controller, but not yet clearly superior in reward robustness.",
            "The contrast between lower mean J and lower mean reward should be discussed as an effect of reward clipping and distribution shape.",
            "State averages are very close to baseline, so the benefit is subtle rather than a dramatic process-state shift.",
        ],
    )

    doc.add_heading("Code changes for the next run", level=1)
    add_bullets(
        doc,
        [
            "Fix the log header so all 29 columns are named from the first row.",
            "Log applied action and next action separately to avoid reward/action timing ambiguity.",
            "Add time/day/episode context to every row for better run-window analysis.",
            "Review reward scaling: Qec may be used more than baseline without consistently improving reward.",
            "Save best checkpoints by rolling J or rolling reward, not only the latest model.",
        ],
    )

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    note_run = note.add_run(
        "One-sentence thesis takeaway: CTRL-1 is stable and cost-promising, but reward robustness and logging need refinement before claiming robust superiority over manual control."
    )
    note_run.bold = True
    note_run.font.color.rgb = RGBColor(31, 78, 121)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
